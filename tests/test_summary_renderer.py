from __future__ import annotations

from pathlib import Path

from docx import Document

from summary_renderer import load_summary, parse_docx_summary, parse_text_summary


def build_variant_summary(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("Grant Summary: AVF 18.005 — Example")
    doc.add_paragraph("1 Narrative summary")
    doc.add_paragraph("Readable narrative.")
    doc.add_paragraph("2 Key facts")
    facts = doc.add_table(rows=1, cols=2)
    facts.rows[0].cells[0].text = "Item"
    facts.rows[0].cells[1].text = "Detail"
    for key, value in (
        ("Grant ID", "AVF 18.005"),
        ("Academic year", "2017-2018 (activities 2018-19)"),
    ):
        cells = facts.add_row().cells
        cells[0].text, cells[1].text = key, value

    doc.add_paragraph("3 CES evidence assessment")
    assessment = doc.add_table(rows=1, cols=3)
    assessment.rows[0].cells[0].text = "Criterion"
    assessment.rows[0].cells[1].text = "Result"
    assessment.rows[0].cells[2].text = "Reasoning"
    cells = assessment.add_row().cells
    cells[0].text = "Community need"
    cells[1].text = "Yes"
    cells[2].text = "Demonstrated by partner requests."
    cells = assessment.add_row().cells
    cells[0].text = "Confidence"
    cells[1].text = "0.60"
    cells[2].text = "Internal score."

    doc.add_paragraph("4 Evidence quotes")
    doc.add_paragraph(
        "Student learning\n"
        "Quotes: “Example quote.”\n"
        "Source file: Final Report.pdf\n"
        "Why it matters: Demonstrates learning."
    )
    doc.add_paragraph("5 Missing or uncertain information")
    doc.add_paragraph("One missing item.")
    doc.save(path)


def test_reads_docx_tables_and_reasoning_column(tmp_path: Path):
    path = tmp_path / "summary.docx"
    build_variant_summary(path)
    data = parse_docx_summary(path)

    assert data.narrative == ["Readable narrative."]
    assert ("Grant ID", "AVF 18.005") in data.facts
    assert ("Academic year", "2017-2018") in data.facts
    assert data.assessment == [
        ("Community need", "Yes", "Demonstrated by partner requests.")
    ]
    assert data.evidence[0]["source"] == "Final Report.pdf"
    assert data.missing == ["One missing item."]


def test_text_fallback_always_preserves_third_assessment_field():
    data = parse_text_summary(
        """Grant Summary: AVF 18.005 — Example
1. Narrative summary
Narrative.
2. Key facts
Item | Detail
Grant ID | AVF 18.005
3. CEL evidence assessment
Criterion | Met? | Explanation
Community need | Yes | Demonstrated.
Overall classification | CEL |
Confidence | 0.60 | Internal score.
4. Evidence quotes
5. Missing or uncertain information
"""
    )
    assert data.assessment == [
        ("Community need", "Yes", "Demonstrated."),
        ("Overall classification", "CEL", ""),
    ]
    assert all(len(row) == 3 for row in data.assessment)


def test_docx_is_preferred_over_stale_flattened_text(tmp_path: Path):
    path = tmp_path / "summary.docx"
    build_variant_summary(path)
    data = load_summary("Old unreadable text only", docx_path=path)
    assert data.facts
    assert data.assessment[0][2] == "Demonstrated by partner requests."


def test_summary_detection_works_with_uuid_filename(tmp_path: Path):
    path = tmp_path / "e341010a-94fa-4bb7-9796-2b5ef5fea04c.docx"
    build_variant_summary(path)
    from summary_renderer import summary_document_score
    assert summary_document_score(path) >= 20
    data = parse_docx_summary(path)
    assert data.assessment[0][2] == "Demonstrated by partner requests."
