from __future__ import annotations

import html
import re
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

try:
    import streamlit as st
except ModuleNotFoundError:  # Parser tests can run without Streamlit.
    st = None

SECTION_RE = re.compile(r"^\s*(?:section\s+)?([1-5])\s*[.):-]?\s*(.+?)\s*$", re.IGNORECASE)
LIST_RE = re.compile(r"^\s*(?:[•●▪◦‣–—-]|\[\[BULLET\]\])\s*")
EVIDENCE_PREFIX_RE = re.compile(
    r"^(?:evidence\s+)?(quotes?|source(?:\s+(?:file|document))?|why\s+(?:it|this)\s+matters|relevance)\s*:\s*(.*)$",
    re.IGNORECASE,
)
NOISE = {"top of form", "bottom of form"}
STATUS_PATTERN = re.compile(
    r"^\s*(not\s+cel|partially|partial|uncertain|unclear|yes|no|cel|met|not\s+met)\b[\s:;,.\-–—]*(.*)$",
    re.IGNORECASE,
)


@dataclass
class SummaryData:
    title: str = ""
    narrative: list[str] = field(default_factory=list)
    facts: list[tuple[str, str]] = field(default_factory=list)
    assessment: list[tuple[str, str, str]] = field(default_factory=list)
    assessment_notes: list[tuple[str, str]] = field(default_factory=list)
    evidence: list[dict[str, str]] = field(default_factory=list)
    missing: list[str] = field(default_factory=list)
    general: list[str] = field(default_factory=list)


def _clean(value: str) -> str:
    return re.sub(r"\s+", " ", (value or "").replace("\u00a0", " ")).strip()


def _strip_list(value: str) -> str:
    return LIST_RE.sub("", _clean(value)).strip()


def _norm(value: str) -> str:
    value = _strip_list(value).casefold()
    value = re.sub(r"\([^)]*\)", " ", value)
    value = re.sub(r"[^a-z0-9]+", " ", value)
    return re.sub(r"\s+", " ", value).strip()


def _paragraph_is_list(paragraph: Paragraph) -> bool:
    properties = paragraph._p.pPr
    return bool(properties is not None and properties.numPr is not None)


def _section_kind(value: str) -> str | None:
    normalized = _norm(value)
    if normalized.startswith("narrative") and "summary" in normalized:
        return "narrative"
    if normalized in {"key facts", "key project facts", "project key facts"} or (
        normalized.startswith("key ") and ("fact" in normalized or "information" in normalized)
    ):
        return "facts"
    if normalized in {"cel assessment", "ces assessment", "cel evidence assessment", "ces evidence assessment"}:
        return "assessment"
    if normalized.startswith(("cel ", "ces ")) and (
        "assessment" in normalized or "classification" in normalized
    ):
        return "assessment"
    if normalized.startswith("evidence ") and ("assessment" in normalized or "classification" in normalized):
        return "assessment"
    if normalized in {"evidence", "evidence quote", "evidence quotes", "supporting evidence"}:
        return "evidence"
    if normalized.startswith("missing") or normalized in {"uncertain information", "information gaps", "information gap"}:
        return "missing"
    return None


def _heading_kind(line: str) -> str | None:
    match = SECTION_RE.match(line)
    if match:
        return _section_kind(match.group(2))
    if len(line) > 100 or ":" in line:
        return None
    return _section_kind(line)


def _normalize_year(value: str) -> str:
    match = re.search(r"(?<!\d)(\d{4})\s*[-–—]\s*(\d{2}|\d{4})(?!\d)", value)
    if not match:
        return value
    start = int(match.group(1))
    end_text = match.group(2)
    if len(end_text) == 2:
        end = (start // 100) * 100 + int(end_text)
        if end < start:
            end += 100
    else:
        end = int(end_text)
    return f"{start:04d}-{end:04d}"


def _is_fact_header(row: list[str]) -> bool:
    if len(row) < 2:
        return False
    first, second = _norm(row[0]), _norm(row[1])
    return first in {"field", "item", "fact", "attribute", "field item"} and second in {
        "detail", "details", "value", "description"
    }


def _is_assessment_header(row: list[str]) -> bool:
    if len(row) < 2:
        return False
    first, second = _norm(row[0]), _norm(row[1])
    third = _norm(row[2]) if len(row) > 2 else ""
    return (
        first in {"criterion", "criteria", "assessment criterion"}
        and second in {"met", "met yes no", "status", "result", "assessment"}
        and (not third or any(token in third for token in ("explanation", "reason", "reasoning", "evidence", "detail")))
    )


def _append_fact(data: SummaryData, field_name: str, detail: str) -> None:
    field_name, detail = _clean(field_name), _clean(detail)
    if not field_name or not detail:
        return
    if _norm(field_name) == "academic year":
        detail = _normalize_year(detail)
    row = (field_name, detail)
    if not data.facts or data.facts[-1] != row:
        data.facts.append(row)


def _split_status_reasoning(value: str) -> tuple[str, str]:
    value = _clean(value)
    match = STATUS_PATTERN.match(value)
    if match:
        return _clean(match.group(1)), _clean(match.group(2))
    return "", value


def _append_assessment(data: SummaryData, criterion: str, status: str, reasoning: str) -> None:
    criterion, status, reasoning = _clean(criterion), _clean(status), _clean(reasoning)
    normalized = _norm(criterion)

    if not criterion:
        if data.assessment:
            old_criterion, old_status, old_reasoning = data.assessment[-1]
            if not reasoning and status:
                status2, reasoning2 = _split_status_reasoning(status)
                if status2:
                    status, reasoning = status2, reasoning2
                else:
                    reasoning, status = status, ""
            data.assessment[-1] = (
                old_criterion,
                _clean(f"{old_status} {status}"),
                _clean(f"{old_reasoning} {reasoning}"),
            )
        return

    if normalized == "confidence":
        return
    if normalized in {"brief explanation", "assessment note", "summary explanation", "reasoning"}:
        value = reasoning or status
        if value and data.assessment and not data.assessment[-1][2]:
            previous_criterion, previous_status, _ = data.assessment[-1]
            data.assessment[-1] = (previous_criterion, previous_status, value)
        elif value:
            data.assessment_notes.append((criterion, value))
        return

    # Some Word exports collapse the result and explanation into one cell.
    if not reasoning and status:
        split_status, split_reasoning = _split_status_reasoning(status)
        if split_status:
            status, reasoning = split_status, split_reasoning

    row = (criterion, status, reasoning)
    if not data.assessment or data.assessment[-1] != row:
        data.assessment.append(row)


def _parse_table(data: SummaryData, rows: list[list[str]], current_kind: str | None) -> None:
    rows = [[_clean(cell) for cell in row] for row in rows if any(_clean(cell) for cell in row)]
    if not rows:
        return

    kind = current_kind
    if _is_fact_header(rows[0]):
        kind = "facts"
        rows = rows[1:]
    elif _is_assessment_header(rows[0]):
        kind = "assessment"
        rows = rows[1:]
    elif len(rows[0]) >= 3 and _norm(rows[0][0]) in {"criterion", "criteria"}:
        kind = "assessment"
        rows = rows[1:]
    elif len(rows[0]) >= 2 and _norm(rows[0][0]) in {"field", "item", "attribute"}:
        kind = "facts"
        rows = rows[1:]

    if kind == "facts":
        for row in rows:
            if len(row) >= 2:
                _append_fact(data, row[0], " ".join(value for value in row[1:] if value))
        return

    if kind == "assessment":
        for row in rows:
            criterion = row[0] if row else ""
            status = row[1] if len(row) > 1 else ""
            reasoning = " ".join(value for value in row[2:] if value)
            _append_assessment(data, criterion, status, reasoning)
        return

    # Header-independent recovery for slightly altered templates.
    for row in rows:
        if len(row) >= 3:
            _append_assessment(data, row[0], row[1], " ".join(row[2:]))
        elif len(row) >= 2:
            _append_fact(data, row[0], row[1])


def _parse_evidence_lines(lines: list[str]) -> list[dict[str, str]]:
    entries: list[dict[str, str]] = []
    current: dict[str, str] | None = None
    active: str | None = None

    def finish() -> None:
        nonlocal current, active
        if current and any(current.get(key) for key in ("quote", "source", "why")):
            entries.append(current)
        current = None
        active = None

    for raw in lines:
        line = _strip_list(raw)
        if not line:
            continue
        match = EVIDENCE_PREFIX_RE.match(line)
        if match:
            if current is None:
                current = {"label": "Evidence"}
            prefix, value = _norm(match.group(1)), _clean(match.group(2))
            active = "quote" if prefix.startswith("quote") else "source" if prefix.startswith("source") else "why"
            if value:
                current[active] = value
            continue
        if current and any(current.get(key) for key in ("quote", "source", "why")):
            finish()
            current = {"label": line}
        elif current is None:
            current = {"label": line}
        elif active:
            current[active] = _clean(f"{current.get(active, '')} {line}")
        else:
            current["label"] = _clean(f"{current.get('label', '')} {line}")
    finish()
    return entries


def summary_document_score(path: Path) -> int:
    """Score any DOCX by its content, not only its filename.

    This lets the app recognize summaries renamed to UUIDs or other generic names.
    """
    if path.suffix.lower() != ".docx" or not path.exists():
        return -1
    score = 8 if "summary" in path.name.casefold() else 0
    try:
        document = Document(path)
        paragraph_text = [_clean(p.text) for p in document.paragraphs if _clean(p.text)]
        first = paragraph_text[0] if paragraph_text else ""
        joined = "\n".join(paragraph_text[:20])
        if _norm(first).startswith("grant summary"):
            score += 30
        kinds = {_heading_kind(line) for line in paragraph_text[:30]}
        score += 7 * len({"narrative", "facts", "assessment", "evidence", "missing"} & kinds)
        for table in document.tables:
            rows = [[_clean(cell.text) for cell in row.cells] for row in table.rows[:2]]
            if rows and _is_fact_header(rows[0]):
                score += 12
            if rows and _is_assessment_header(rows[0]):
                score += 12
        if "grant id" in _norm(joined):
            score += 3
    except Exception:
        return score - 20
    return score


def parse_docx_summary(path: Path) -> SummaryData:
    document = Document(path)
    data = SummaryData()
    current_kind: str | None = None
    evidence_lines: list[str] = []

    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            paragraph = Paragraph(child, document)
            raw_text = paragraph.text or ""
            text = _clean(raw_text)
            if not text or _norm(text) in NOISE:
                continue
            if not data.title and _norm(text).startswith("grant summary"):
                data.title = text
                continue
            kind = _heading_kind(text)
            if kind:
                current_kind = kind
                continue
            if _paragraph_is_list(paragraph):
                text = f"• {text}"
            if current_kind == "narrative":
                data.narrative.append(text)
            elif current_kind == "evidence":
                evidence_lines.extend(
                    cleaned for part in raw_text.splitlines()
                    if (cleaned := _clean(part))
                )
            elif current_kind == "missing":
                data.missing.append(_strip_list(text))
            elif current_kind == "facts":
                if ":" in text:
                    key, value = text.split(":", 1)
                    _append_fact(data, key, value)
                else:
                    data.general.append(text)
            elif current_kind == "assessment":
                parts = [_clean(part) for part in re.split(r"\s*[|\t]\s*", text) if _clean(part)]
                if len(parts) >= 2:
                    _append_assessment(data, parts[0], parts[1], " ".join(parts[2:]))
                else:
                    data.general.append(text)
            else:
                data.general.append(text)
        elif isinstance(child, CT_Tbl):
            table = Table(child, document)
            rows = [[cell.text for cell in row.cells] for row in table.rows]
            _parse_table(data, rows, current_kind)

    data.evidence = _parse_evidence_lines(evidence_lines)
    return data


def parse_text_summary(text: str) -> SummaryData:
    data = SummaryData()
    current_kind: str | None = None
    evidence_lines: list[str] = []
    table_mode: str | None = None

    for raw_line in (text or "").splitlines():
        line = _clean(raw_line)
        if not line or _norm(line) in NOISE:
            continue
        if not data.title and _norm(line).startswith("grant summary"):
            data.title = line
            continue
        kind = _heading_kind(line)
        if kind:
            current_kind = kind
            table_mode = None
            continue

        if "|" in line:
            parts = [_clean(part) for part in line.split("|")]
            if _is_fact_header(parts):
                table_mode = "facts"
                continue
            if _is_assessment_header(parts):
                table_mode = "assessment"
                continue
            effective = table_mode or current_kind
            if effective == "facts" and len(parts) >= 2:
                _append_fact(data, parts[0], " ".join(value for value in parts[1:] if value))
                continue
            if effective == "assessment" and len(parts) >= 2:
                _append_assessment(data, parts[0], parts[1], " ".join(value for value in parts[2:] if value))
                continue

        if current_kind == "narrative":
            data.narrative.append(line)
        elif current_kind == "evidence":
            evidence_lines.append(line)
        elif current_kind == "missing":
            data.missing.append(_strip_list(line))
        elif current_kind == "facts" and ":" in line:
            key, value = line.split(":", 1)
            _append_fact(data, key, value)
        elif current_kind == "assessment":
            data.general.append(line)
        else:
            data.general.append(line)

    data.evidence = _parse_evidence_lines(evidence_lines)
    return data


def _merge(primary: SummaryData, fallback: SummaryData) -> SummaryData:
    if not primary.title:
        primary.title = fallback.title
    if not primary.narrative:
        primary.narrative = fallback.narrative
    if not primary.facts:
        primary.facts = fallback.facts
    if not primary.assessment:
        primary.assessment = fallback.assessment
    if not primary.assessment_notes:
        primary.assessment_notes = fallback.assessment_notes
    if not primary.evidence:
        primary.evidence = fallback.evidence
    if not primary.missing:
        primary.missing = fallback.missing
    if not primary.general:
        primary.general = fallback.general
    return primary


def load_summary(text: str, docx_path: Path | None = None) -> SummaryData:
    fallback = parse_text_summary(text)
    if docx_path and docx_path.exists():
        try:
            return _merge(parse_docx_summary(docx_path), fallback)
        except Exception:
            return fallback
    return fallback


def _styles() -> None:
    st.markdown(
        """
        <style>
        .grant-summary-copy p { line-height: 1.68; margin: 0 0 1rem 0; }
        .grant-table-wrap { width: 100%; overflow-x: auto; margin: .25rem 0 1rem 0; }
        .grant-table { width: 100%; border-collapse: collapse; table-layout: fixed; }
        .grant-table th { text-align: left; padding: .65rem .75rem; border-bottom: 2px solid rgba(128,128,128,.35); }
        .grant-table td { vertical-align: top; padding: .65rem .75rem; border-bottom: 1px solid rgba(128,128,128,.22); line-height: 1.5; overflow-wrap: anywhere; }
        .grant-table.facts th:first-child, .grant-table.facts td:first-child { width: 28%; font-weight: 650; }
        .grant-table.ces th:nth-child(1), .grant-table.ces td:nth-child(1) { width: 23%; font-weight: 600; }
        .grant-table.ces th:nth-child(2), .grant-table.ces td:nth-child(2) { width: 12%; }
        .grant-table.ces th:nth-child(3), .grant-table.ces td:nth-child(3) { width: 65%; }
        .grant-evidence { border: 1px solid rgba(128,128,128,.25); border-radius: .6rem; padding: .9rem 1rem; margin: 0 0 .85rem 0; }
        .grant-evidence-title { font-weight: 700; margin-bottom: .45rem; }
        .grant-evidence-quote { border-left: 3px solid rgba(128,128,128,.45); padding-left: .8rem; line-height: 1.58; margin-bottom: .55rem; }
        .grant-evidence-source { opacity: .75; font-size: .9rem; margin-top: .4rem; }
        .grant-list { line-height: 1.65; }
        @media (max-width: 760px) {
          .grant-table.ces, .grant-table.ces thead, .grant-table.ces tbody, .grant-table.ces th, .grant-table.ces td, .grant-table.ces tr { display: block; width: 100% !important; }
          .grant-table.ces thead { display: none; }
          .grant-table.ces tr { border: 1px solid rgba(128,128,128,.25); border-radius: .55rem; margin-bottom: .75rem; padding: .3rem .5rem; }
          .grant-table.ces td { border: 0; padding: .35rem .45rem; }
          .grant-table.ces td::before { display: block; font-weight: 700; opacity: .72; font-size: .78rem; text-transform: uppercase; margin-bottom: .15rem; }
          .grant-table.ces td:nth-child(1)::before { content: "Criterion"; }
          .grant-table.ces td:nth-child(2)::before { content: "Result"; }
          .grant-table.ces td:nth-child(3)::before { content: "Reasoning"; }
        }
        </style>
        """,
        unsafe_allow_html=True,
    )


def _paragraph(value: str) -> None:
    st.markdown(
        f'<div class="grant-summary-copy"><p>{html.escape(_strip_list(value))}</p></div>',
        unsafe_allow_html=True,
    )


def _table(rows: list[tuple[str, ...]], headers: tuple[str, ...], css_class: str) -> None:
    if not rows:
        st.caption("No structured information was included in this section.")
        return
    head = "".join(f"<th>{html.escape(header)}</th>" for header in headers)
    body = "".join(
        "<tr>" + "".join(f"<td>{html.escape(value or '—')}</td>" for value in row) + "</tr>"
        for row in rows
    )
    st.markdown(
        f'<div class="grant-table-wrap"><table class="grant-table {css_class}"><thead><tr>{head}</tr></thead><tbody>{body}</tbody></table></div>',
        unsafe_allow_html=True,
    )


def _bullets(values: list[str]) -> None:
    items = "".join(f"<li>{html.escape(_strip_list(value))}</li>" for value in values if _strip_list(value))
    if items:
        st.markdown(f'<ul class="grant-list">{items}</ul>', unsafe_allow_html=True)


def render_summary(text: str, docx_path: Path | None = None) -> None:
    if st is None:
        raise RuntimeError("Streamlit is required to render summaries.")
    data = load_summary(text, docx_path)
    _styles()

    if data.title:
        st.caption(data.title)

    st.markdown("### 1. Narrative summary")
    narrative = data.narrative or data.general
    if narrative:
        for paragraph in narrative:
            _paragraph(paragraph)
    else:
        st.caption("No narrative summary was found.")
    st.divider()

    st.markdown("### 2. Key facts")
    _table(data.facts, ("Field", "Detail"), "facts")
    st.divider()

    st.markdown("### 3. CES evidence assessment")
    # Every displayed assessment row always has a third reasoning column.
    assessment_rows = [(criterion, result, reasoning or "—") for criterion, result, reasoning in data.assessment]
    _table(assessment_rows, ("Criterion", "Result", "Reasoning"), "ces")
    for label, value in data.assessment_notes:
        st.info(f"{label}: {value}")
    st.divider()

    st.markdown("### 4. Evidence quotes")
    if data.evidence:
        for item in data.evidence:
            pieces = [f'<div class="grant-evidence-title">{html.escape(item.get("label", "Evidence"))}</div>']
            if item.get("quote"):
                pieces.append(f'<div class="grant-evidence-quote">{html.escape(item["quote"])}</div>')
            if item.get("why"):
                pieces.append(f'<div>{html.escape(item["why"])}</div>')
            if item.get("source"):
                pieces.append(f'<div class="grant-evidence-source">Source: {html.escape(item["source"])}</div>')
            st.markdown(f'<div class="grant-evidence">{"".join(pieces)}</div>', unsafe_allow_html=True)
    else:
        st.caption("No evidence quotes were included.")
    st.divider()

    st.markdown("### 5. Missing or uncertain information")
    if data.missing:
        _bullets(data.missing)
    else:
        st.caption("No missing-information notes were included.")
